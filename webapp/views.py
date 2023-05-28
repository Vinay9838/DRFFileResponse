from io import BytesIO

import pandas as pd
from docx import Document
from rest_framework.views import APIView
from rest_framework.response import Response  
from django.http import FileResponse
from rest_framework import status
from drf_spectacular.utils import extend_schema,OpenApiResponse,OpenApiParameter

from .models import Employee


# Create your views here.

class EmployeeViewSet(APIView):

    def get_excel(self,queryset):
        file_name = 'Sample.xlsx'
        byte_buffer = BytesIO()
        df = pd.DataFrame.from_dict(queryset.values())
        #df2 = pd.DataFrame.from_dict(employee2)    
        writer = pd.ExcelWriter(byte_buffer,engine='xlsxwriter')
        df.to_excel(writer,sheet_name='Employee',index=False)
        #df2.to_excel(writer,sheet_name='Research and Development',index=False)
        writer.close()
        return byte_buffer,file_name
    

    def get_document(self,queryset):
        df = pd.DataFrame.from_dict(queryset[:10].values('id','first_name','last_name','gender','job_title'))
        file_name = "Sample.docx"
        document = Document()
        document.add_heading("Employee List Sample")
        document.add_paragraph("This is sample document for employee list detail")
        table = document.add_table(rows=1, cols=len(df.columns.values.tolist()))
        table.autofit = True
        hdr_cells = table.rows[0].cells
        for i,col in enumerate(df.columns.values.tolist()):
            hdr_cells[i].text = col
        
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i,col in enumerate(df.columns.values.tolist()):
                row_cells[i].text = str(row[col])
        document.add_page_break()
        byte_buffer = BytesIO()
        document.save(byte_buffer)  # save your memory stream
        return byte_buffer,file_name
        
    def get_text(self):
        file_name = "Sample.txt"
        sample_text = """
        Lorem ipsum dolor sit amet, consectetur adipiscing elit. Sed ac congue lorem. Nudrerit rutrum ipsum tincidunt gravida. Duis tempor dapibus libero. Nulla in erat magna. Praesent libero lectus, congue id turpis quis, consequat placerat magna. Vestibulum ante ipsum primis in faucibus orci luctus et ultrices posuere cubilia curae; Proin iaculis mauris quis augue lacinia, a feugiat nisi fermentum. Pellentesque felis erat, consectetur ac tellus a, sollicitudin semper dolor. Vivamus facilisis massa ex, vitae accumsan urna convallis quis. Nam blandit lorem arcu, at finibus ex molestie vel. Sed a elit urna. Morbi massa metus, placerat nec quam a, vulputate pretium velit. Praesent pharetra quam egestas, porta dui aliquet, bibendum est.
        """
        byte_buffer = BytesIO(sample_text.encode("utf-8"))
        return byte_buffer,file_name

    @extend_schema(
        summary='Send Binary response',
        parameters=[
            OpenApiParameter(name='res_type',type=str,enum=['text','excel','docs'],required=True,location=OpenApiParameter.PATH)
        ],
        responses={
            200: OpenApiResponse(description='Binary Response'),
            404: OpenApiResponse(description='Resource not found')
        }
    )
    def post(self,request,res_type): 
        employee = Employee.objects.all()
        if res_type == 'excel':
            byte_buffer,file_name = self.get_excel(employee)
        elif res_type == 'text':
            byte_buffer,file_name = self.get_text()
        elif res_type == 'docs':
            byte_buffer,file_name = self.get_document(employee)
        else:
            return Response({"error":"Invalid res_type"},status=status.HTTP_400_BAD_REQUEST)
        byte_buffer.seek(0)
        return FileResponse(byte_buffer,filename=file_name,as_attachment=True)

